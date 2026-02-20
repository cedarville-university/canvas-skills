#!/usr/bin/env python3
"""Extract a Course Alignment Grid (DOCX) into Canvas buildRequest JSON.

Modes:
- deterministic: rule-based parser only
- llm: OpenAI-only extraction
- auto: deterministic first, then OpenAI fallback on parse failure/low confidence
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

try:
    from dotenv import load_dotenv
except Exception:  # pragma: no cover
    load_dotenv = None

try:
    import jsonschema
except Exception:  # pragma: no cover
    jsonschema = None


DEFAULTS = {
    "course_id": -1,
    "start_date": "2025-08-26 00:00:00",
    "end_date": "2025-12-15 23:59:59",
    "default_due_day": 6,
    "default_discussion_due_day": 3,
    "default_last_day": 4,
    "build_type": 2,
    "overview_page_template": "Module 1: Overview",
    "discussion_template": "Group Discussion: [Title Here]",
    "assignment_template": "Individual Assignment: [Title Here]",
    "newquiz_template": "New Quiz: [Title Here]",
    "classicquiz_template": "Classic Quiz: [Title Here]",
    "llm_model": "gpt-5-mini",
}


def maybe_load_dotenv() -> None:
    if load_dotenv is None:
        return
    # Let explicit environment values win.
    load_dotenv(override=False)


def prompt_text(prompt: str, default: str = "", allow_empty: bool = True) -> str:
    while True:
        suffix = f" [{default}]" if default else ""
        try:
            value = input(f"{prompt}{suffix}: ").strip()
        except EOFError:
            return default if default else ""
        if value:
            return value
        if default:
            return default
        if allow_empty:
            return ""
        print("A value is required.")


def prompt_int(prompt: str, default: int = 0, allow_empty: bool = True) -> int:
    while True:
        raw = prompt_text(prompt, str(default) if default else "", allow_empty=allow_empty)
        if raw == "" and allow_empty:
            return default
        if re.fullmatch(r"-?\d+", raw):
            return int(raw)
        print("Enter an integer value.")


def parse_pipe_list(value: str) -> list[str]:
    return [normalize_text(part) for part in value.split("|") if normalize_text(part)]


def normalize_text(value: str) -> str:
    value = value.replace("\u00a0", " ")
    value = value.replace("\u2013", "-")
    value = value.replace("\u2014", "-")
    value = value.replace("\uf0b7", " ")
    value = value.strip()
    value = re.sub(r"^[\-*\u2022]+\s*", "", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def strip_trailing_objective_noise(value: str) -> str:
    value = normalize_text(value)
    value = re.sub(r"\s*\([^)]*\)\s*$", "", value).strip()
    value = re.sub(r"\s+\d+(?:\.\d+)*\s*$", "", value).strip()
    value = value.rstrip(".,;:")
    return value


def strip_module_notes(value: str) -> str:
    value = normalize_text(value)
    value = re.sub(r"\([^)]*\)", "", value)
    value = re.sub(r"\s+:", ":", value)
    value = re.sub(r"\s+", " ", value).strip()
    if value.upper().startswith("MODULE "):
        value = "Module " + value[7:]
    return value


def get_value_after_label(text: str, label: str) -> str:
    match = re.search(rf"{re.escape(label)}\s*:\s*(.+)", text, flags=re.IGNORECASE)
    if not match:
        return ""
    return normalize_text(match.group(1))


def split_items(text: str) -> list[str]:
    raw = text.replace("\r", "\n")
    items: list[str] = []
    for line in raw.split("\n"):
        line = normalize_text(line)
        if not line:
            continue
        parts = [normalize_text(part) for part in line.split("|")]
        for part in parts:
            if part:
                items.append(part)
    return items


def replace_inline_links(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", r'<a href="\2">\1</a>', text)
    text = re.sub(r"(https?://[^\s)]+)", r'<a href="\1">\1</a>', text)
    return text


def parse_file_marker(item: str) -> tuple[str, int] | None:
    match = re.search(r"(?i)\bfile\s*:\s*(\d+)\b", item)
    if not match:
        return None
    file_id = int(match.group(1))
    name = re.sub(r"(?i)\bfile\s*:\s*\d+\b", "", item).strip(" -:")
    name = normalize_text(name) if name else ""
    if not name:
        name = f"File {file_id}"
    return name, file_id


def parse_module_objectives(text: str) -> list[str]:
    return [strip_trailing_objective_noise(item) for item in split_items(text) if item]


@dataclass
class Counters:
    quiz: int = 0
    discussion: int = 0
    assignment: int = 0
    page: int = 0


def classify_assignment(name: str) -> tuple[str, str]:
    lower = name.lower()

    if "exam" in lower:
        if "classic" in lower:
            return "classic quiz", name
        return "quiz", name

    if "discussion" in lower:
        return "discussion", name

    if "classic quiz" in lower:
        cleaned = re.sub(r"(?i)\bclassic\b", "", name)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        return "classic quiz", cleaned

    if "quiz" in lower:
        return "quiz", name

    return "assignment", name


def assignment_id(assignment_type: str, counters: Counters) -> str:
    if assignment_type in {"quiz", "classic quiz"}:
        counters.quiz += 1
        return f"q{counters.quiz}"
    if assignment_type == "discussion":
        counters.discussion += 1
        return f"d{counters.discussion}"
    counters.assignment += 1
    return f"a{counters.assignment}"


def extract_explicit_assignment_id(value: str) -> tuple[str, str | int | None]:
    """
    Extract explicit assignment id markers and return cleaned name + id.

    Supported markers:
    - id:123
    - id=q7
    - (id: a12)
    - [id: q3]
    - {id=abc_01}
    """
    pattern = re.compile(r"(?i)[\[\(\{]?\s*id\s*[:=]\s*([a-z0-9_-]+)\s*[\]\)\}]?")
    match = pattern.search(value)
    if not match:
        return value, None

    raw_id = match.group(1).strip()
    explicit_id: str | int
    if raw_id.isdigit():
        explicit_id = int(raw_id)
    else:
        explicit_id = raw_id

    cleaned = (value[: match.start()] + " " + value[match.end() :]).strip()
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -:;,")
    return cleaned, explicit_id


def build_modules(table, content_course_id: str) -> list[dict[str, Any]]:
    row_count = len(table.rows)
    col_count = len(table.columns)
    if col_count not in {4, 5}:
        raise ValueError(f"Unsupported CAG table format: expected 4 or 5 columns, found {col_count}")

    if col_count == 4:
        overview_idx = None
        objectives_idx = 1
        assessments_idx = 2
        content_idx = 3
    else:
        overview_idx = 1
        objectives_idx = 2
        assessments_idx = 3
        content_idx = 4

    modules: list[dict[str, Any]] = []
    counters = Counters()
    module_number = 1
    i = 1

    while i < row_count:
        row = table.rows[i]
        cells = [cell.text or "" for cell in row.cells]
        module_name = strip_module_notes(cells[0])

        if not module_name:
            i += 1
            continue

        detail_row = row
        if module_name.lower().startswith("module") and i + 1 < row_count:
            next_cells = [cell.text or "" for cell in table.rows[i + 1].cells]
            next_first = strip_module_notes(next_cells[0])
            if not next_first.lower().startswith("module"):
                detail_row = table.rows[i + 1]
                i += 1

        detail_cells = [cell.text or "" for cell in detail_row.cells]
        overview = normalize_text(detail_cells[overview_idx]) if overview_idx is not None else ""
        module_objectives = parse_module_objectives(detail_cells[objectives_idx])
        assessment_items = split_items(detail_cells[assessments_idx])
        content_items = split_items(detail_cells[content_idx])

        assignments: list[dict[str, Any]] = []
        for item in assessment_items:
            normalized_item = normalize_text(item)
            cleaned_item, explicit_id = extract_explicit_assignment_id(normalized_item)
            assign_type, assign_name = classify_assignment(cleaned_item)
            selected_id = explicit_id if explicit_id is not None else assignment_id(assign_type, counters)
            assignments.append(
                {
                    "id": selected_id,
                    "name": assign_name,
                    "type": assign_type,
                }
            )

        pages: list[dict[str, Any]] = []
        files: list[dict[str, Any]] = []
        content: list[str] = []
        for item in content_items:
            normalized = normalize_text(item)

            if normalized == "#new_page" or "(new_page)" in normalized.lower():
                title = re.sub(r"(?i)\(new_page\)", "", normalized)
                title = title.replace("#new_page", "").strip(" -:")
                title = normalize_text(title)
                counters.page += 1
                pages.append({"id": f"p{counters.page}", "title": title})
                content.append(f'<a href="#new_page">{title}</a>')
                continue

            file_marker = parse_file_marker(normalized)
            if file_marker:
                file_name, file_id = file_marker
                files.append({"id": file_id, "name": file_name})
                content.append(
                    '<a class="instructure_file_link instructure_scribd_file inline_disabled" '
                    f'href="/courses/{content_course_id}/files/{file_id}?wrap=1" target="_blank" '
                    f'rel="noopener noreferrer">{file_name}</a>'
                )
                continue

            content.append(replace_inline_links(normalized))

        modules.append(
            {
                "id": module_number,
                "name": module_name,
                "number": module_number,
                "position": module_number + 3,
                "overview": overview,
                "objectives": module_objectives,
                "assessments": [],
                "assignments": assignments,
                "content": content,
                "pages": pages,
                "files": files,
            }
        )
        module_number += 1
        i += 1

    return modules


def section_paragraphs(doc: Document, start_heading: str, end_heading: str | None = None) -> list[str]:
    capture = False
    out: list[str] = []
    start_key = start_heading.lower()
    end_key = end_heading.lower() if end_heading else None

    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue

        lower = text.lower()
        if lower == start_key:
            capture = True
            continue
        if capture and end_key and lower == end_key:
            break
        if capture:
            out.append(text)
    return out


def extract_course_data(doc: Document, content_course_id: str) -> dict[str, Any]:
    paragraph_text = "\n".join(para.text for para in doc.paragraphs)

    course_code = get_value_after_label(paragraph_text, "Course Code")
    course_name = get_value_after_label(paragraph_text, "Course title")
    instructor_name = get_value_after_label(paragraph_text, "Instructor")
    credits_raw = get_value_after_label(paragraph_text, "Credit")
    year_raw = get_value_after_label(paragraph_text, "Year")
    term = get_value_after_label(paragraph_text, "Term")
    start_at = get_value_after_label(paragraph_text, "Start_at")
    end_at = get_value_after_label(paragraph_text, "End_at")

    textbooks = [normalize_text(item) for item in section_paragraphs(doc, "Textbook:", "Course policy")]
    policy_candidates = section_paragraphs(doc, "Course policy", "Course Overview")
    course_policy_parts: list[str] = []
    for para in policy_candidates:
        if para.endswith(":"):
            continue
        if "." not in para and ":" not in para and len(para.split()) <= 6:
            continue
        course_policy_parts.append(para)

    description_section = section_paragraphs(doc, "Course Overview", "Current Course Objectives")
    description = ""
    for para in description_section:
        if para.lower() == "course description":
            continue
        description = para
        break

    objective_items = section_paragraphs(doc, "Current Course Objectives", "Course alignment grid")
    objectives = [strip_trailing_objective_noise(item) for item in objective_items if item]

    if not doc.tables:
        raise ValueError("No tables found in DOCX. Expected a CAG table.")
    modules = build_modules(doc.tables[0], content_course_id)

    return {
        "course_code": course_code,
        "course_name": course_name,
        "description": description,
        "instructor": [{"name": instructor_name, "email": ""}] if instructor_name else [],
        "year": int(year_raw) if year_raw.isdigit() else 0,
        "term": term,
        "start_at": start_at,
        "end_at": end_at,
        "credits": int(credits_raw) if credits_raw.isdigit() else 0,
        "objectives": objectives,
        "textbooks": textbooks,
        "course_policy": " ".join(course_policy_parts).strip(),
        "modules": modules,
    }


def render_doc_for_llm(doc: Document) -> str:
    lines: list[str] = []
    lines.append("# Paragraphs")
    for idx, para in enumerate(doc.paragraphs, 1):
        text = normalize_text(para.text)
        if text:
            lines.append(f"P{idx}: {text}")

    lines.append("")
    lines.append("# Tables")
    for t_index, table in enumerate(doc.tables, 1):
        lines.append(f"Table {t_index}: rows={len(table.rows)} cols={len(table.columns)}")
        for r_index, row in enumerate(table.rows, 1):
            cells = [normalize_text(cell.text) for cell in row.cells]
            lines.append(f"R{r_index}: " + " || ".join(cells))
        lines.append("")
    return "\n".join(lines)


def extract_json_object(raw: str) -> dict[str, Any]:
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        text = text.strip()

    try:
        data = json.loads(text)
        if isinstance(data, dict):
            return data
    except Exception:
        pass

    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("LLM output did not contain a JSON object")

    data = json.loads(text[start : end + 1])
    if not isinstance(data, dict):
        raise ValueError("LLM output JSON was not an object")
    return data


def response_to_text(response: Any) -> str:
    output_text = getattr(response, "output_text", None)
    if output_text:
        return output_text

    parts: list[str] = []
    output = getattr(response, "output", None)
    if isinstance(output, list):
        for item in output:
            content = getattr(item, "content", None)
            if isinstance(content, list):
                for chunk in content:
                    text = getattr(chunk, "text", None)
                    if text:
                        parts.append(text)
    return "\n".join(parts).strip()


def load_instruction_text(path: Path | None) -> str:
    if path and path.exists():
        return path.read_text()

    default_path = Path(__file__).resolve().parent.parent / "references" / "cag-build-request.md"
    if default_path.exists():
        return default_path.read_text()

    return (
        "Extract course metadata and modules from the CAG document. "
        "Output only JSON with the required course fields and module content."
    )


def normalize_module_shape(module: Any, index: int) -> dict[str, Any]:
    if not isinstance(module, dict):
        module = {}

    return {
        "id": int(module.get("id", index)),
        "name": str(module.get("name", "")),
        "number": int(module.get("number", index)),
        "position": int(module.get("position", index + 3)),
        "overview": str(module.get("overview", "")),
        "objectives": list(module.get("objectives", [])) if isinstance(module.get("objectives", []), list) else [],
        "assessments": list(module.get("assessments", [])) if isinstance(module.get("assessments", []), list) else [],
        "assignments": list(module.get("assignments", [])) if isinstance(module.get("assignments", []), list) else [],
        "content": list(module.get("content", [])) if isinstance(module.get("content", []), list) else [],
        "pages": list(module.get("pages", [])) if isinstance(module.get("pages", []), list) else [],
        "files": list(module.get("files", [])) if isinstance(module.get("files", []), list) else [],
    }


def normalize_course_shape(course: Any) -> dict[str, Any]:
    if not isinstance(course, dict):
        course = {}

    modules = course.get("modules", [])
    if not isinstance(modules, list):
        modules = []

    normalized_modules = [normalize_module_shape(module, idx + 1) for idx, module in enumerate(modules)]

    instructor = course.get("instructor", [])
    if not isinstance(instructor, list):
        instructor = []

    objectives = course.get("objectives", [])
    textbooks = course.get("textbooks", [])

    return {
        "course_code": str(course.get("course_code", "")),
        "course_name": str(course.get("course_name", "")),
        "description": str(course.get("description", "")),
        "instructor": instructor,
        "year": int(course.get("year", 0) or 0),
        "term": str(course.get("term", "")),
        "start_at": str(course.get("start_at", "")),
        "end_at": str(course.get("end_at", "")),
        "credits": int(course.get("credits", 0) or 0),
        "objectives": objectives if isinstance(objectives, list) else [],
        "textbooks": textbooks if isinstance(textbooks, list) else [],
        "course_policy": str(course.get("course_policy", "")),
        "modules": normalized_modules,
    }


def low_confidence(course: dict[str, Any]) -> bool:
    modules = course.get("modules", [])
    if not modules:
        return True
    if not course.get("course_code") and not course.get("course_name"):
        return True

    empty_modules = 0
    for module in modules:
        if not module.get("name"):
            empty_modules += 1
            continue
        if not module.get("objectives") and not module.get("assignments") and not module.get("content"):
            empty_modules += 1

    return empty_modules > len(modules) // 2


def prompt_missing_course_fields(course: dict[str, Any], interactive: bool) -> dict[str, Any]:
    if not interactive:
        return course

    if not course.get("course_code"):
        course["course_code"] = prompt_text("Missing course_code")
    if not course.get("course_name"):
        course["course_name"] = prompt_text("Missing course_name")
    if not course.get("description"):
        course["description"] = prompt_text("Missing description")

    instructor = course.get("instructor") if isinstance(course.get("instructor"), list) else []
    if not instructor:
        name = prompt_text("Missing instructor name", allow_empty=True)
        if name:
            instructor = [{"name": name, "email": ""}]
    else:
        first = instructor[0] if isinstance(instructor[0], dict) else {"name": "", "email": ""}
        if not first.get("name"):
            first["name"] = prompt_text("Missing instructor name", allow_empty=True)
        if "email" not in first:
            first["email"] = ""
        instructor[0] = first
    course["instructor"] = instructor

    if not course.get("year"):
        course["year"] = prompt_int("Missing year", default=0, allow_empty=True)
    if not course.get("term"):
        course["term"] = prompt_text("Missing term", allow_empty=True)
    if not course.get("start_at"):
        course["start_at"] = prompt_text("Missing start_at (YYYY-MM-DD)", allow_empty=True)
    if not course.get("end_at"):
        course["end_at"] = prompt_text("Missing end_at (YYYY-MM-DD)", allow_empty=True)
    if not course.get("credits"):
        course["credits"] = prompt_int("Missing credits", default=0, allow_empty=True)

    if not course.get("objectives"):
        raw = prompt_text(
            "Missing objectives. Enter objectives separated by |",
            allow_empty=True,
        )
        course["objectives"] = parse_pipe_list(raw)

    if not course.get("textbooks"):
        raw = prompt_text(
            "Missing textbooks. Enter textbooks separated by |",
            allow_empty=True,
        )
        course["textbooks"] = parse_pipe_list(raw)

    if not course.get("course_policy"):
        course["course_policy"] = prompt_text("Missing course_policy", allow_empty=True)

    return course


def prompt_missing_build_fields(args: argparse.Namespace, course: dict[str, Any]) -> argparse.Namespace:
    if not args.interactive:
        return args

    if args.course_id == -1:
        args.course_id = prompt_int("Missing course_id (Canvas course ID)", default=-1, allow_empty=True)

    if args.start_date is None and not course.get("start_at"):
        args.start_date = prompt_text(
            "Missing start_date (YYYY-MM-DD HH:MM:SS)",
            default=DEFAULTS["start_date"],
            allow_empty=True,
        )

    if args.end_date is None and not course.get("end_at"):
        args.end_date = prompt_text(
            "Missing end_date (YYYY-MM-DD HH:MM:SS)",
            default=DEFAULTS["end_date"],
            allow_empty=True,
        )

    if args.course_id != -1:
        marker = "/courses/{courseid}/files/"
        replacement = f"/courses/{args.course_id}/files/"
        for module in course.get("modules", []):
            content = module.get("content", [])
            if not isinstance(content, list):
                continue
            module["content"] = [
                item.replace(marker, replacement) if isinstance(item, str) else item for item in content
            ]

    return args


def extract_course_data_with_llm(
    doc: Document,
    model: str,
    instructions_text: str,
    schema_text: str,
    api_key_env: str,
) -> dict[str, Any]:
    try:
        from openai import OpenAI
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "openai package is required for --mode llm/auto fallback. Install with: python3 -m pip install openai"
        ) from exc

    api_key = os.environ.get(api_key_env)
    if not api_key:
        raise RuntimeError(f"Missing {api_key_env} for LLM parsing mode")

    schema_hint = ""
    if schema_text:
        try:
            schema_obj = json.loads(schema_text)
            course_schema = schema_obj.get("properties", {}).get("course", {})
            schema_hint = json.dumps(course_schema or schema_obj, indent=2)
        except Exception:
            schema_hint = schema_text

    system_prompt = (
        "You are a strict extraction engine. Return exactly one JSON object. "
        "No markdown, no comments, no code fences."
    )
    user_prompt = (
        "Extract a JSON object for the 'course' payload from the provided DOCX content.\n"
        "Follow these extraction instructions:\n"
        f"{instructions_text}\n\n"
        "Course schema target (use all keys; unknown values must be empty strings/lists/0):\n"
        f"{schema_hint}\n\n"
        "DOCX content:\n"
        f"{render_doc_for_llm(doc)}\n"
    )

    client = OpenAI(api_key=api_key)

    last_error: Exception | None = None
    # First try Responses API.
    try:
        response = client.responses.create(
            model=model,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        data = extract_json_object(response_to_text(response))
        course = data.get("course") if "course" in data else data
        return normalize_course_shape(course)
    except Exception as exc:
        last_error = exc

    # Fallback to Chat Completions JSON mode.
    try:
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            response_format={"type": "json_object"},
        )
        raw = completion.choices[0].message.content or "{}"
        data = extract_json_object(raw)
        course = data.get("course") if "course" in data else data
        return normalize_course_shape(course)
    except Exception as exc:
        raise RuntimeError(f"LLM parsing failed: {exc}; responses error: {last_error}") from exc


def build_payload(course: dict[str, Any], args: argparse.Namespace) -> dict[str, Any]:
    start_date = args.start_date or (
        f"{course['start_at']} 00:00:00" if course.get("start_at") else DEFAULTS["start_date"]
    )
    end_date = args.end_date or (
        f"{course['end_at']} 23:59:59" if course.get("end_at") else DEFAULTS["end_date"]
    )

    return {
        "course_id": args.course_id,
        "start_date": start_date,
        "end_date": end_date,
        "default_due_day": args.default_due_day,
        "default_discussion_due_day": args.default_discussion_due_day,
        "default_last_day": args.default_last_day,
        "build_type": args.build_type,
        "overview_page_template": args.overview_page_template,
        "discussion_template": args.discussion_template,
        "assignment_template": args.assignment_template,
        "newquiz_template": args.newquiz_template,
        "classicquiz_template": args.classicquiz_template,
        "course": course,
    }


def validate_with_schema(payload: dict[str, Any], schema_path: Path) -> None:
    if jsonschema is None:
        raise RuntimeError("jsonschema is not installed. Install with: python3 -m pip install jsonschema")
    schema = json.loads(schema_path.read_text())
    jsonschema.validate(instance=payload, schema=schema)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract a CAG DOCX into buildRequest JSON payload."
    )
    parser.add_argument("--input-docx", required=True, type=Path, help="Path to source .docx file")
    parser.add_argument("--output-json", required=False, type=Path, help="Output .json path")
    parser.add_argument("--schema", required=False, type=Path, help="Optional JSON schema path")
    parser.add_argument(
        "--interactive",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Prompt for missing fields before generating final payload (default: on; use --no-interactive to disable)",
    )
    parser.add_argument(
        "--mode",
        choices=["auto", "deterministic", "llm"],
        default="auto",
        help="auto tries deterministic first and falls back to LLM when needed",
    )
    parser.add_argument("--llm-model", default=DEFAULTS["llm_model"])
    parser.add_argument("--instructions", type=Path, default=None, help="Optional extraction instruction markdown")
    parser.add_argument("--api-key-env", default="OPENAI_API_KEY", help="Environment variable for OpenAI API key")

    parser.add_argument("--course-id", type=int, default=DEFAULTS["course_id"])
    parser.add_argument("--start-date", default=None)
    parser.add_argument("--end-date", default=None)
    parser.add_argument("--default-due-day", type=int, default=DEFAULTS["default_due_day"])
    parser.add_argument(
        "--default-discussion-due-day",
        type=int,
        default=DEFAULTS["default_discussion_due_day"],
    )
    parser.add_argument("--default-last-day", type=int, default=DEFAULTS["default_last_day"])
    parser.add_argument("--build-type", type=int, default=DEFAULTS["build_type"])
    parser.add_argument(
        "--overview-page-template",
        default=DEFAULTS["overview_page_template"],
    )
    parser.add_argument("--discussion-template", default=DEFAULTS["discussion_template"])
    parser.add_argument("--assignment-template", default=DEFAULTS["assignment_template"])
    parser.add_argument("--newquiz-template", default=DEFAULTS["newquiz_template"])
    parser.add_argument("--classicquiz-template", default=DEFAULTS["classicquiz_template"])
    return parser.parse_args()


def run_extraction(args: argparse.Namespace, doc: Document, content_course_id: str) -> dict[str, Any]:
    deterministic_error: Exception | None = None

    if args.mode in {"deterministic", "auto"}:
        try:
            course = extract_course_data(doc, content_course_id=content_course_id)
            if args.mode == "auto" and low_confidence(course):
                raise ValueError("deterministic output confidence too low")
            return normalize_course_shape(course)
        except Exception as exc:
            deterministic_error = exc
            if args.mode == "deterministic":
                raise

    if args.mode in {"llm", "auto"}:
        schema_text = args.schema.read_text() if args.schema else ""
        instructions_text = load_instruction_text(args.instructions)
        try:
            return extract_course_data_with_llm(
                doc=doc,
                model=args.llm_model,
                instructions_text=instructions_text,
                schema_text=schema_text,
                api_key_env=args.api_key_env,
            )
        except Exception as exc:
            if deterministic_error:
                raise RuntimeError(
                    f"deterministic parse failed: {deterministic_error}; llm parse failed: {exc}"
                ) from exc
            raise

    raise RuntimeError(f"Unsupported mode: {args.mode}")


def main() -> int:
    args = parse_args()
    maybe_load_dotenv()

    if not args.input_docx.exists():
        print(f"Input DOCX not found: {args.input_docx}", file=sys.stderr)
        return 1

    if args.interactive and args.course_id == -1:
        args.course_id = prompt_int("Missing course_id (Canvas course ID)", default=-1, allow_empty=True)

    doc = Document(str(args.input_docx))
    content_course_id = str(args.course_id if args.course_id != -1 else "{courseid}")

    try:
        course = run_extraction(args, doc, content_course_id)
        course = prompt_missing_course_fields(course, args.interactive)
        args = prompt_missing_build_fields(args, course)
        payload = build_payload(course, args)

        if args.schema:
            validate_with_schema(payload, args.schema)
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    output_text = json.dumps(payload, indent=2, ensure_ascii=True)
    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(output_text + "\n")
        print(f"Wrote {args.output_json}")
    else:
        print(output_text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
